# -*- coding: utf-8 -*-
"""Generates a docx file of a homework that can be given to a
japanese elementary(?) student studying english.
"""
from sqlalchemy import create_engine, MetaData, Table, desc
from sqlalchemy.orm import sessionmaker
from sqlalchemy.engine.url import URL as dburl
from sqlalchemy.sql.expression import func, bindparam
from random import seed, randint, shuffle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor, Inches


db = dburl(
    drivername='mysql+pymysql',
    host='',  # change host name
    port='3306',  # change port if necessary
    username='',  # change username
    password='',  # change password
    database='homework',  # database name where vocab table is located
    query={'charset': 'utf8'})
engine = create_engine(db, encoding='utf-8', echo=False)
maker = sessionmaker(autocommit=True, autoflush=True, bind=engine)
session = maker()
metadata = MetaData()

font_segoe_print = 'Segoe Print'
font_ms_mincho = 'MS Mincho'
font_times_new_roman = 'Times New Roman'
title_font_size = 14
star_font_size = 8
table_content_font_size = 10.5
title_font_color = RGBColor(0xff, 0x33, 0xcc)
table_header_font_color = RGBColor(0x46, 0x74, 0xb5)
table_headers = ['英語', '日本語', '品詞', '用例／頻出の連語']
cell_widths = [1.51, 2.25, 0.88, 2.01]
cell_widths_question = [2.14, 2.27, 0.71, 1.52]
the_star = '☆彡'
blank = '(                    )'
blank2 = '(                  )'
en_text = '{0}    {1}'
space_before = 2
space_after = 2

def addFromList(file):
    with open(file) as f:
        contents = f.readlines()

    types = Table('types', metadata, autoload=True, autoload_with=engine)
    types_list = session.query(types.columns.id, types.columns.jp).all()
    types_dict = {}
    [types_dict.update({row.jp: row.id}) for row in types_list]

    contents = [x.strip() for x in contents]
    contents = [x.split(',') for x in contents]
    new_rows = []
    for item in contents:
        if len(item) == 0:
            continue
        type_str = item[1].strip()
        type_split = type_str.split('|')
        type_str = ''
        for i, a_type in enumerate(type_split):
            print(a_type)
            type_fixed = a_type
            if a_type == '前':
                a_type += '置詞'
            elif a_type == '形':
                a_type += '容詞'
            elif a_type == '代':
                a_type += '名詞'
            elif a_type == '接':
                a_type += '続詞'
            elif a_type == '疑':
                a_type += '問詞'
            elif a_type == '間':
                a_type += '投詞'
            elif a_type == '助':
                a_type += '動詞'
            elif a_type == '関':
                a_type += '係代名詞'
            elif a_type == 'n/a':
                pass
            else:
                a_type += '詞'
            print(item[0].strip(), a_type)
            type_str += str(types_dict[a_type])
            if i != len(type_split) - 1:
                type_str += ','
        new_rows.append({'en': item[0].strip(),
                         'jp': item[2].strip(),
                         'type': type_str,
                         'example': item[3].strip()})

    vocab = Table('vocab', metadata, autoload=True, autoload_with=engine)
    result = session.execute(vocab.insert(), new_rows)
    inserted = result.last_inserted_params()
    return True if len(inserted) == len(new_rows) else False

def generate(count, test_number):
    """Generates a docx file with <i>count<i> randomly picked words
    and <i>test_number<i> as the nth homework.
    """
    types = Table('types', metadata, autoload=True, autoload_with=engine)
    types_list = session.query(types.columns.id, types.columns.jp).all()
    types_dict = {}
    [types_dict.update({str(row.id): row.jp}) for row in types_list]

    vocab = Table('vocab', metadata, autoload=True, autoload_with=engine)
    c = vocab.columns
    result = session.query(c.id, c.useCount, c.en, c.jp, c.type, c.example) \
        .filter(c.useCount == 0) \
        .order_by(func.rand()) \
        .limit(count).all()
    fetched_ids = [{'_id': row.id,
                    'count': row.useCount + 1} for row in result]
    print(fetched_ids)

    items = []
    for row in result:
        item = {}
        item.update({'en': row.en,
                     'jp': row.jp,
                     'example': row.example,
                     'pick': randint(0, 1)})
        the_type = row.type
        type_str = ''
        if ',' in the_type:
            type_split = the_type.split(',')
            for i, a_type in enumerate(type_split):
                type_str += types_dict[a_type]
                if i != len(type_split) - 1:
                    type_str += '|'
        else:
            type_str = '' if the_type == 'n/a' else types_dict[the_type]
        item.update({'type': type_str})
        print(item)
        items.append(item)

    document = Document()
    # answer title
    for i, title in enumerate(['Answers for',
                               'Vocabulary test No.{}'.format(test_number)]):
        title_paragraph = document.add_paragraph()
        run = title_paragraph.add_run(title)
        font = run.font
        font.color.rgb = title_font_color
        font.size = Pt(title_font_size)
        font.name = font_segoe_print
        font.bold = True
        font.underline = True
        title_format = title_paragraph.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_before = Pt(space_before)
        title_format.space_after = Pt(space_after)
        title_format.line_spacing = Pt(0)
        if i == 1:
            star_run = title_paragraph.add_run(the_star)
            star_font = star_run.font
            star_font.color.rgb = title_font_color
            star_font.size = Pt(star_font_size)
            star_font.name = font_ms_mincho
            star_font.bold = True
    document.add_paragraph('')
    # answers table
    table = document.add_table(rows=count + 1, cols=4)
    for i, row in enumerate(table.rows):
        cells = row.cells
        # headers
        if i == 0:
            for j, header in enumerate(table_headers):
                cell = cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_paragraph = cell.paragraphs[0]
                cell_format = cell_paragraph.paragraph_format
                cell_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_format.space_before = Pt(space_before)
                cell_format.space_after = Pt(space_after)
                cell_format.line_spacing = Pt(0)
                cell_run = cell_paragraph.add_run(header)
                cell_font = cell_run.font
                cell_font.color.rgb = table_header_font_color
                cell_font.size = Pt(table_content_font_size)
                cell_font.bold = True
                cell_font.underline = True
                cell.width = Inches(cell_widths[j])
        else:
            item = items[i - 1]
            for k, column in enumerate(['en', 'jp', 'type', 'example']):
                col_item = item[column]
                if k == 0:
                    cells[k].text = en_text.format(str(i), col_item)
                else:
                    col_split = col_item.split('|')
                    for j, col_word in enumerate(col_split):
                        if j == 0:
                            cells[k].text = col_word
                        else:
                            cells[k].add_paragraph(col_word)
                col_cell_paragraphs = cells[k].paragraphs
                for cell_paragraph in col_cell_paragraphs:
                    cell_format = cell_paragraph.paragraph_format
                    cell_format.space_before = Pt(space_before)
                    cell_format.space_after = Pt(space_after)
                    cell_format.line_spacing = Pt(0)
                    for cell_run in cell_paragraph.runs:
                        cell_font = cell_run.font
                        if k in (0, 3):
                            cell_font.name = font_times_new_roman
                        else:
                            cell_font.name = font_ms_mincho
                        cell_font.size = Pt(table_content_font_size)
                cells[k].width = Inches(cell_widths[k])
    # source part
    document.add_paragraph('')
    source = document.add_paragraph('（大阪版中学校で学ぶ英単語集'
                                    '（約1,300語）より）')
    source_format = source.paragraph_format
    source_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source_format.space_before = Pt(space_before)
    source_format.space_after = Pt(space_after)
    source_format.line_spacing = Pt(0)

    document.add_page_break()

    # questions title
    title_paragraph = document.add_paragraph()
    run = title_paragraph.add_run('Vocabulary test No.{}'.format(test_number))
    font = run.font
    font.color.rgb = title_font_color
    font.size = Pt(title_font_size)
    font.name = font_segoe_print
    font.bold = True
    font.underline = True
    title_format = title_paragraph.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format.space_before = Pt(space_before)
    title_format.space_after = Pt(space_after)
    title_format.line_spacing = Pt(0)
    star_run = title_paragraph.add_run(the_star)
    star_font = star_run.font
    star_font.color.rgb = title_font_color
    star_font.size = Pt(star_font_size)
    star_font.name = font_ms_mincho
    star_font.bold = True
    document.add_paragraph('')
    # questions table
    tableq = document.add_table(rows=count + 1, cols=4)
    for i, row in enumerate(tableq.rows):
        cells = row.cells
        #headers
        if i == 0:
            for j, header in enumerate(table_headers):
                cell = cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_paragraph = cell.paragraphs[0]
                cell_format = cell_paragraph.paragraph_format
                cell_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_format.space_before = Pt(space_before)
                cell_format.space_after = Pt(space_after)
                cell_format.line_spacing = Pt(0)
                cell_run = cell_paragraph.add_run(header)
                cell_font = cell_run.font
                cell_font.color.rgb = table_header_font_color
                cell_font.size = Pt(table_content_font_size)
                cell_font.bold = True
                cell_font.underline = True
                cell.width = Inches(cell_widths_question[j])
        else:
            item = items[i - 1]
            pick = item['pick']
            for k, column in enumerate(['en', 'jp', 'type', 'example']):
                col_item = item[column]
                if k == 0:
                    if pick == 0:
                        cells[k].text = en_text.format(str(i), col_item)
                    else:
                        en_paragraph = cells[k].paragraphs[0]
                        number_run = en_paragraph.add_run(str(i))
                        number_font = number_run.font
                        number_font.name = font_times_new_roman
                        number_font.size = Pt(table_content_font_size)
                        if i > 9:
                            blank_str = ' ' + blank2
                        else:
                            blank_str = '  ' + blank2
                        blank_run = en_paragraph.add_run(blank_str)
                        blank_font = blank_run.font
                        blank_font.name = font_ms_mincho
                        blank_font.size = Pt(table_content_font_size)
                elif k in (1, 2):
                    col_split = col_item.split('|')
                    for j, col_word in enumerate(col_split):
                        if j == 0:
                            cells[k].text = col_word
                        else:
                            cells[k].add_paragraph(col_word)
                    if k == 1 and pick != 1:
                        cells[k].text = blank
                elif k == 3:
                    cells[k].text = ''
                col_cell_paragraphs = cells[k].paragraphs
                for cell_paragraph in col_cell_paragraphs:
                    cell_format = cell_paragraph.paragraph_format
                    cell_format.space_before = Pt(space_before)
                    cell_format.space_after = Pt(space_after)
                    cell_format.line_spacing = Pt(0)
                    for cell_run in cell_paragraph.runs:
                        cell_font = cell_run.font
                        if k in (1, 2):
                            cell_font.name = font_ms_mincho
                        cell_font.size = Pt(table_content_font_size)
                cells[k].width = Inches(cell_widths_question[k])

    document.save('test{}.docx'.format(test_number))
    update_list(fetched_ids)

def generate_using(word_list, test_number):
    types = Table('types', metadata, autoload=True, autoload_with=engine)
    types_list = session.query(types.columns.id, types.columns.jp).all()
    types_dict = {}
    [types_dict.update({str(row.id): row.jp}) for row in types_list]

    vocab = Table('vocab', metadata, autoload=True, autoload_with=engine)
    items = []
    fetched_ids = []
    for word in word_list:
        result = session.query(vocab.columns.id,
                               vocab.columns.useCount,
                               vocab.columns.en,
                               vocab.columns.jp,
                               vocab.columns.type,
                               vocab.columns.example) \
            .filter(vocab.columns.en == word).first()
        if result is not None:
            item = {}
            item.update({'en': result.en,
                         'jp': result.jp,
                         'example': result.example,
                         'pick': randint(0, 1)})
            fetched_ids.append({'_id': result.id,
                                'count': result.useCount + 1})
            the_type = result.type
            type_str = ''
            if ',' in the_type:
                type_split = the_type.split(',')
                for i, a_type in enumerate(type_split):
                    type_str += types_dict[a_type]
                    if i != len(type_split) - 1:
                        type_str += '|'
            else:
                type_str = '' if the_type == 'n/a' else types_dict[the_type]
            item.update({'type': type_str})
            items.append(item)
            print(items)
    count = len(items)

    document = Document()
    # answer title
    for i, title in enumerate(['Answers for',
                               'Vocabulary test No.{}'.format(test_number)]):
        title_paragraph = document.add_paragraph()
        run = title_paragraph.add_run(title)
        font = run.font
        font.color.rgb = title_font_color
        font.size = Pt(title_font_size)
        font.name = font_segoe_print
        font.bold = True
        font.underline = True
        title_format = title_paragraph.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format.space_before = Pt(space_before)
        title_format.space_after = Pt(space_after)
        title_format.line_spacing = Pt(0)
        if i == 1:
            star_run = title_paragraph.add_run(the_star)
            star_font = star_run.font
            star_font.color.rgb = title_font_color
            star_font.size = Pt(star_font_size)
            star_font.name = font_ms_mincho
            star_font.bold = True
    document.add_paragraph('')
    # answers table
    table = document.add_table(rows=count + 1, cols=4)
    for i, row in enumerate(table.rows):
        cells = row.cells
        # headers
        if i == 0:
            for j, header in enumerate(table_headers):
                cell = cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_paragraph = cell.paragraphs[0]
                cell_format = cell_paragraph.paragraph_format
                cell_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_format.space_before = Pt(space_before)
                cell_format.space_after = Pt(space_after)
                cell_format.line_spacing = Pt(0)
                cell_run = cell_paragraph.add_run(header)
                cell_font = cell_run.font
                cell_font.color.rgb = table_header_font_color
                cell_font.size = Pt(table_content_font_size)
                cell_font.bold = True
                cell_font.underline = True
                cell.width = Inches(cell_widths[j])
        else:
            item = items[i - 1]
            for k, column in enumerate(['en', 'jp', 'type', 'example']):
                col_item = item[column]
                if k == 0:
                    cells[k].text = en_text.format(str(i), col_item)
                else:
                    col_split = col_item.split('|')
                    for j, col_word in enumerate(col_split):
                        if j == 0:
                            cells[k].text = col_word
                        else:
                            cells[k].add_paragraph(col_word)
                col_cell_paragraphs = cells[k].paragraphs
                for cell_paragraph in col_cell_paragraphs:
                    cell_format = cell_paragraph.paragraph_format
                    cell_format.space_before = Pt(space_before)
                    cell_format.space_after = Pt(space_after)
                    cell_format.line_spacing = Pt(0)
                    for cell_run in cell_paragraph.runs:
                        cell_font = cell_run.font
                        if k in (0, 3):
                            cell_font.name = font_times_new_roman
                        else:
                            cell_font.name = font_ms_mincho
                        cell_font.size = Pt(table_content_font_size)
                cells[k].width = Inches(cell_widths[k])
    # source part
    document.add_paragraph('')
    source = document.add_paragraph('（大阪版中学校で学ぶ英単語集'
                                    '（約1,300語）より）')
    source_format = source.paragraph_format
    source_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source_format.space_before = Pt(space_before)
    source_format.space_after = Pt(space_after)
    source_format.line_spacing = Pt(0)

    document.add_page_break()

    # questions title
    title_paragraph = document.add_paragraph()
    run = title_paragraph.add_run('Vocabulary test No.{}'.format(test_number))
    font = run.font
    font.color.rgb = title_font_color
    font.size = Pt(title_font_size)
    font.name = font_segoe_print
    font.bold = True
    font.underline = True
    title_format = title_paragraph.paragraph_format
    title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format.space_before = Pt(space_before)
    title_format.space_after = Pt(space_after)
    title_format.line_spacing = Pt(0)
    star_run = title_paragraph.add_run(the_star)
    star_font = star_run.font
    star_font.color.rgb = title_font_color
    star_font.size = Pt(star_font_size)
    star_font.name = font_ms_mincho
    star_font.bold = True
    document.add_paragraph('')
    # questions table
    tableq = document.add_table(rows=count + 1, cols=4)
    for i, row in enumerate(tableq.rows):
        cells = row.cells
        #headers
        if i == 0:
            for j, header in enumerate(table_headers):
                cell = cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell_paragraph = cell.paragraphs[0]
                cell_format = cell_paragraph.paragraph_format
                cell_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_format.space_before = Pt(space_before)
                cell_format.space_after = Pt(space_after)
                cell_format.line_spacing = Pt(0)
                cell_run = cell_paragraph.add_run(header)
                cell_font = cell_run.font
                cell_font.color.rgb = table_header_font_color
                cell_font.size = Pt(table_content_font_size)
                cell_font.bold = True
                cell_font.underline = True
                cell.width = Inches(cell_widths_question[j])
        else:
            item = items[i - 1]
            pick = item['pick']
            for k, column in enumerate(['en', 'jp', 'type', 'example']):
                col_item = item[column]
                if k == 0:
                    if pick == 0:
                        cells[k].text = en_text.format(str(i), col_item)
                    else:
                        en_paragraph = cells[k].paragraphs[0]
                        number_run = en_paragraph.add_run(str(i))
                        number_font = number_run.font
                        number_font.name = font_times_new_roman
                        number_font.size = Pt(table_content_font_size)
                        if i > 9:
                            blank_str = ' ' + blank2
                        else:
                            blank_str = '  ' + blank2
                        blank_run = en_paragraph.add_run(blank_str)
                        blank_font = blank_run.font
                        blank_font.name = font_ms_mincho
                        blank_font.size = Pt(table_content_font_size)
                elif k in (1, 2):
                    col_split = col_item.split('|')
                    for j, col_word in enumerate(col_split):
                        if j == 0:
                            cells[k].text = col_word
                        else:
                            cells[k].add_paragraph(col_word)
                    if k == 1 and pick != 1:
                        cells[k].text = blank
                elif k == 3:
                    cells[k].text = ''
                col_cell_paragraphs = cells[k].paragraphs
                for cell_paragraph in col_cell_paragraphs:
                    cell_format = cell_paragraph.paragraph_format
                    cell_format.space_before = Pt(space_before)
                    cell_format.space_after = Pt(space_after)
                    cell_format.line_spacing = Pt(0)
                    for cell_run in cell_paragraph.runs:
                        cell_font = cell_run.font
                        if k in (1, 2):
                            cell_font.name = font_ms_mincho
                        cell_font.size = Pt(table_content_font_size)
                cells[k].width = Inches(cell_widths_question[k])

    document.save('test{}.docx'.format(test_number))
    update_list(fetched_ids)

def update_list(used_ids):
    vocab = Table('vocab', metadata, autoload=True, autoload_with=engine)
    statement = vocab.update() \
        .where(vocab.columns.id == bindparam('_id')) \
        .values({'useCount': bindparam('count')})
    result = session.execute(statement, used_ids)
    print(result)

def reset_list():
    vocab = Table('vocab', metadata, autoload=True, autoload_with=engine)
    result = session.execute(vocab.update().values(useCount=0))
    print(result)
